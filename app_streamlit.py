import streamlit as st
import os, re, zipfile, tempfile
import fitz
import docx
import numpy as np
import pandas as pd
from PIL import Image
import imagehash
from io import BytesIO

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from core.file_operations import extract_text_from_pdf


#text utilities


def split_sentences(text):
    if not isinstance(text, str):
        return []
    return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 20]

def normalize(s):
    return re.sub(r"\s+", " ", s.lower().strip())

def remove_template_text(text, template_text):
    if not template_text:
        return text
    template_set = set(normalize(s) for s in split_sentences(template_text))
    return " ".join(s for s in split_sentences(text) if normalize(s) not in template_set)

def highlight_html(text, scores, threshold):
    out = []
    for s in split_sentences(text):
        score = scores.get(s, 0)
        esc = s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        if score >= threshold:
            out.append(
                f"<span style='background:#ffeb3b;font-weight:bold'>{esc}</span>"
                f"<small> ({round(score*100,1)}%)</small>"
            )
        else:
            out.append(esc)
    return "<br><br>".join(out)

#image extraction


def extract_images_from_pdf(path):
    images = []
    doc = fitz.open(path)

    for page in doc:
        for img in page.get_images(full=True):
            pix = fitz.Pixmap(doc, img[0])
            if pix.n > 3:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            try:
                pil = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append({
                    "hash": imagehash.phash(pil),
                    "image": pil
                })
            except:
                pass
    return images

def extract_images_from_docx(path):
    images = []
    d = docx.Document(path)

    for rel in d.part._rels.values():
        if "image" in rel.reltype:
            img_bytes = rel.target_part.blob
            try:
                pil = Image.open(BytesIO(img_bytes)).convert("RGB")
                images.append({
                    "hash": imagehash.phash(pil),
                    "image": pil
                })
            except:
                pass
    return images

#image filtering and matching

def exclude_template_images(images, template_hashes, threshold=8):
    if not template_hashes:
        return images
    clean = []
    for img in images:
        if all(img["hash"] - th > threshold for th in template_hashes):
            clean.append(img)
    return clean

def match_images(images1, images2, threshold=8):
    matches = []
    used = set()

    for i, a in enumerate(images1):
        for j, b in enumerate(images2):
            if j in used:
                continue
            if a["hash"] - b["hash"] <= threshold:
                matches.append((a["image"], b["image"]))
                used.add(j)
                break
    return matches

#text similarity

def text_similarity(t1, t2):
    s1, s2 = split_sentences(t1), split_sentences(t2)
    if not s1 or not s2:
        return {}, {}, 0.0

    vec = TfidfVectorizer(stop_words="english")
    tfidf = vec.fit_transform(s1 + s2)

    sim = cosine_similarity(tfidf[:len(s1)], tfidf[len(s1):])
    scores1 = {s1[i]: float(sim[i].max()) for i in range(len(s1))}
    scores2 = {s2[j]: float(sim[:, j].max()) for j in range(len(s2))}
    return scores1, scores2, float(np.mean(list(scores1.values())))

#streamlit ui

st.set_page_config("Plagiarism Checker", layout="wide")
st.title("Plagiarism Checker (PDF + DOCX)")

st.sidebar.header("Upload")
file_type = st.sidebar.selectbox("Document Type", ["PDF", "DOCX"])
zips = st.sidebar.file_uploader("Upload ZIP files", type=["zip"], accept_multiple_files=True)
template_file = st.sidebar.file_uploader("Template File (exclude text & images)",
    type=["pdf"] if file_type == "PDF" else ["docx"]
)
threshold = st.sidebar.slider("Text highlight threshold (%)", 30, 90, 55) / 100

tmp = tempfile.mkdtemp()

#template

template_text = ""
template_image_hashes = []

if template_file:
    tpath = os.path.join(tmp, template_file.name)
    with open(tpath, "wb") as f:
        f.write(template_file.read())

    if file_type == "PDF":
        template_text = extract_text_from_pdf(tpath)
        imgs = extract_images_from_pdf(tpath)
    else:
        d = docx.Document(tpath)
        template_text = " ".join(p.text for p in d.paragraphs)
        imgs = extract_images_from_docx(tpath)

    template_image_hashes = [i["hash"] for i in imgs]

#zipfile extraction

for z in zips:
    zp = os.path.join(tmp, z.name)
    with open(zp, "wb") as f:
        f.write(z.read())
    with zipfile.ZipFile(zp) as zr:
        zr.extractall(tmp)

#load documents

docs = []

for root, _, files in os.walk(tmp):
    for f in files:
        if file_type == "PDF" and f.lower().endswith(".pdf"):
            path = os.path.join(root, f)
            text = extract_text_from_pdf(path)
            images = extract_images_from_pdf(path)

        elif file_type == "DOCX" and f.lower().endswith(".docx"):
            path = os.path.join(root, f)
            d = docx.Document(path)
            text = " ".join(p.text for p in d.paragraphs)
            images = extract_images_from_docx(path)
        else:
            continue

        if template_file and f == template_file.name:
            continue

        text = remove_template_text(text, template_text)
        images = exclude_template_images(images, template_image_hashes)

        docs.append({
            "name": f,
            "text": text,
            "images": images
        })

st.write(f"📂 Documents loaded: {len(docs)}")
if len(docs) < 2:
    st.stop()

#comparison

rows = []

for i in range(len(docs)):
    for j in range(i+1, len(docs)):
        A, B = docs[i], docs[j]

        s1, s2, text_sim = text_similarity(A["text"], B["text"])
        img_matches = match_images(A["images"], B["images"])

        img_sim = len(img_matches) / max(1, min(len(A["images"]), len(B["images"])))
        combined = 0.6 * text_sim + 0.4 * img_sim

        rows.append({
            "File 1": A["name"],
            "File 2": B["name"],
            "Text %": round(text_sim * 100, 2),
            "Image %": round(img_sim * 100, 2),
            "Combined %": round(combined * 100, 2),
            "scores1": s1,
            "scores2": s2,
            "images": img_matches
        })

df = pd.DataFrame(rows).sort_values("Combined %", ascending=False)

st.subheader("📊 Similarity Results")
st.dataframe(df[["File 1","File 2","Text %","Image %","Combined %"]])

#detailed description of matched text/images

for _, r in df.iterrows():
    st.markdown(f"## {r['File 1']} ↔ {r['File 2']}")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Text – File 1**")
        st.markdown(
            highlight_html(
                next(d["text"] for d in docs if d["name"] == r["File 1"]),
                r["scores1"], threshold
            ),
            unsafe_allow_html=True
        )

    with c2:
        st.markdown("**Text – File 2**")
        st.markdown(
            highlight_html(
                next(d["text"] for d in docs if d["name"] == r["File 2"]),
                r["scores2"], threshold
            ),
            unsafe_allow_html=True
        )

    if r["images"]:
        st.markdown("### 🖼️ Matched Images")
        for a, b in r["images"]:
            i1, i2 = st.columns(2)
            i1.image(a, use_container_width=True)
            i2.image(b, use_container_width=True)
    else:
        st.info("No non-template images matched.")
